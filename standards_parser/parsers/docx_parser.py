"""
DOCX (Word) document parser using python-docx library.
Extracts text and structured data from DOCX files.
"""

import logging
from typing import Optional

from .base import BaseStandardParser

logger = logging.getLogger(__name__)


class DocxParser(BaseStandardParser):
    """
    DOCX parser implementation using python-docx.
    
    Handles text extraction from Word documents and delegates metadata
    parsing to the parent class.
    """
    
    def __init__(self, file_path: str) -> None:
        """
        Initialize DOCX parser.
        
        Args:
            file_path: Path to the DOCX file.
        
        Raises:
            FileNotFoundError: If the DOCX file does not exist.
            ImportError: If python-docx is not installed.
        """
        super().__init__(file_path)
        
        try:
            from docx import Document
            self.Document = Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            )
    
    def extract_text(self) -> str:
        """
        Extract text content from DOCX file.
        
        Reads all paragraphs and tables from the Word document.
        Handles corrupted documents gracefully.
        
        Returns:
            Concatenated text from all paragraphs and tables.
        
        Raises:
            Exception: If DOCX is corrupted or cannot be read.
        """
        try:
            self.logger.info(f"Opening DOCX file: {self.file_path}")
            
            # Open document
            doc = self.Document(str(self.file_path))
            
            full_text = ""
            
            # Extract text from paragraphs
            self.logger.info(f"Extracting text from {len(doc.paragraphs)} paragraphs")
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"
            
            # Extract text from tables if any
            if doc.tables:
                self.logger.info(f"Extracting text from {len(doc.tables)} tables")
                full_text += "\n--- Tables ---\n"
                
                for table_idx, table in enumerate(doc.tables, 1):
                    for row_idx, row in enumerate(table.rows, 1):
                        row_cells = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_cells.append(cell.text.strip())
                        if row_cells:
                            full_text += " | ".join(row_cells) + "\n"
            
            if not full_text.strip():
                raise ValueError("No text content could be extracted from DOCX")
            
            self.logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            self.logger.error(f"DOCX extraction error for {self.file_path}: {str(e)}", exc_info=True)
            raise

